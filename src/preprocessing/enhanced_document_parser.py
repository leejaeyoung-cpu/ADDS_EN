"""
Enhanced Document Parser
Advanced parsing for PDF and DOCX with structure recognition
"""

import re
from pathlib import Path
from typing import Dict, List, Any, Optional
import json

try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None


class EnhancedDocumentParser:
    """Advanced document parser with structure recognition"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse document and extract structured information
        
        Args:
            file_path: Path to document
            
        Returns:
            Structured document data
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            return self.parse_pdf(file_path)
        elif ext == '.docx':
            return self.parse_docx(file_path)
        elif ext == '.txt':
            return self.parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported format: {ext}")
    
    def parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF with table and metadata extraction"""
        
        if pdfplumber is None or PyPDF2 is None:
            raise ImportError("pdfplumber and PyPDF2 required for PDF parsing")
        
        result = {
            "filename": file_path.name,
            "file_type": "PDF",
            "metadata": {},
            "sections": [],
            "tables": [],
            "full_text": "",
            "page_count": 0,
            "entities": {}
        }
        
        # Extract metadata
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                result["page_count"] = len(pdf_reader.pages)
                
                if pdf_reader.metadata:
                    result["metadata"] = {
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                        "creator": pdf_reader.metadata.get('/Creator', ''),
                        "producer": pdf_reader.metadata.get('/Producer', ''),
                        "creation_date": str(pdf_reader.metadata.get('/CreationDate', ''))
                    }
        except Exception as e:
            result["metadata"]["error"] = str(e)
        
        # Extract text and tables
        full_text_parts = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        full_text_parts.append(page_text)
                        
                        # Detect sections (simple heuristic)
                        sections = self._detect_sections(page_text, page_num)
                        result["sections"].extend(sections)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        for table_idx, table in enumerate(tables):
                            result["tables"].append({
                                "page": page_num,
                                "table_index": table_idx,
                                "data": table,
                                "caption": self._extract_table_caption(page_text, table_idx)
                            })
        
        except Exception as e:
            result["parsing_error"] = str(e)
        
        result["full_text"] = "\n\n".join(full_text_parts)
        
        # Extract entities
        result["entities"] = self.extract_entities(result["full_text"])
        
        return result
    
    def parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX with structure recognition"""
        
        if Document is None:
            raise ImportError("python-docx required for DOCX parsing")
        
        result = {
            "filename": file_path.name,
            "file_type": "DOCX",
            "metadata": {},
            "sections": [],
            "tables": [],
            "full_text": "",
            "paragraph_count": 0,
            "entities": {}
        }
        
        try:
            doc = Document(file_path)
            
            # Extract metadata
            core_props = doc.core_properties
            result["metadata"] = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else ""
            }
            
            # Extract paragraphs
            paragraphs = []
            current_section = None
            section_content = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                paragraphs.append(text)
                
                # Detect headings/sections
                if para.style.name.startswith('Heading'):
                    # Save previous section
                    if current_section:
                        result["sections"].append({
                            "title": current_section,
                            "content": "\n".join(section_content),
                            "order": len(result["sections"])
                        })
                    
                    current_section = text
                    section_content = []
                else:
                    section_content.append(text)
            
            # Save last section
            if current_section:
                result["sections"].append({
                    "title": current_section,
                    "content": "\n".join(section_content),
                    "order": len(result["sections"])
                })
            
            result["paragraph_count"] = len(paragraphs)
            result["full_text"] = "\n\n".join(paragraphs)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                
                result["tables"].append({
                    "table_index": table_idx,
                    "data": table_data,
                    "rows": len(table.rows),
                    "cols": len(table.columns)
                })
        
        except Exception as e:
            result["parsing_error"] = str(e)
        
        # Extract entities
        result["entities"] = self.extract_entities(result["full_text"])
        
        return result
    
    def parse_txt(self, file_path: Path) -> Dict[str, Any]:
        """Parse plain text file"""
        
        result = {
            "filename": file_path.name,
            "file_type": "TXT",
            "metadata": {},
            "sections": [],
            "full_text": "",
            "entities": {}
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            result["full_text"] = text
            result["metadata"]["file_size"] = len(text)
            
            # Simple section detection
            sections = self._detect_sections(text, 1)
            result["sections"] = sections
            
            # Extract entities
            result["entities"] = self.extract_entities(text)
        
        except Exception as e:
            result["parsing_error"] = str(e)
        
        return result
    
    def _detect_sections(self, text: str, page_num: int = 1) -> List[Dict]:
        """Simple section detection based on patterns"""
        sections = []
        
        # Pattern for section headers (all caps, numbered, etc.)
        patterns = [
            r'^([A-Z][A-Z\s]+)$',  # All caps headers
            r'^(\d+\.\s+[A-Z][^.]+)',  # Numbered sections
            r'^([IVX]+\.\s+[A-Z][^.]+)',  # Roman numerals
        ]
        
        lines = text.split('\n')
        current_header = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            is_header = False
            for pattern in patterns:
                if re.match(pattern, line):
                    # Save previous section
                    if current_header:
                        sections.append({
                            "title": current_header,
                            "content": "\n".join(current_content),
                            "page": page_num,
                            "order": len(sections)
                        })
                    
                    current_header = line
                    current_content = []
                    is_header = True
                    break
            
            if not is_header and current_header:
                current_content.append(line)
        
        # Save last section
        if current_header:
            sections.append({
                "title": current_header,
                "content": "\n".join(current_content),
                "page": page_num,
                "order": len(sections)
            })
        
        return sections
    
    def _extract_table_caption(self, text: str, table_idx: int) -> str:
        """Extract table caption from nearby text"""
        # Simple heuristic: look for "Table X:" pattern
        pattern = rf'Table\s+{table_idx + 1}[:\.]?\s*([^\n]+)'
        match = re.search(pattern, text, re.IGNORECASE)
        return match.group(1) if match else ""
    
    def extract_entities(self, text: str) -> Dict[str, List[str]]:
        """
        Extract medical/scientific entities
        
        Returns:
            Dictionary of entity types and lists
        """
        entities = {
            "drugs": [],
            "genes": [],
            "proteins": [],
            "diseases": [],
            "dosages": []
        }
        
        # Drug name patterns (simplified)
        drug_patterns = [
            r'\b([A-Z][a-z]+mab)\b',  # Monoclonal antibodies
            r'\b(cisplatin|doxorubicin|paclitaxel|5-FU|irinotecan|gemcitabine|oxaliplatin)\b'
        ]
        
        for pattern in drug_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            entities["drugs"].extend([m for m in matches if m not in entities["drugs"]])
        
        # Gene patterns (all caps, 3-6 letters)
        gene_pattern = r'\b([A-Z]{3,6})\b'
        potential_genes = re.findall(gene_pattern, text)
        # Filter out common words
        common_words = {'THE', 'AND', 'FOR', 'WITH', 'FROM', 'THAT', 'THIS'}
        entities["genes"] = [g for g in set(potential_genes) if g not in common_words][:20]  # Limit
        
        # Dosage patterns
        dosage_pattern = r'(\d+(?:\.\d+)?\s*(?:mg|µg|g|mL|L)(?:/m²)?)'
        entities["dosages"] = list(set(re.findall(dosage_pattern, text)))
        
        return entities
    
    def to_json(self, parsed_data: Dict) -> str:
        """Convert parsed data to JSON string"""
        return json.dumps(parsed_data, indent=2, ensure_ascii=False)
    
    def save_parsed(self, parsed_data: Dict, output_path: str):
        """Save parsed data to JSON file"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(parsed_data, f, indent=2, ensure_ascii=False)
