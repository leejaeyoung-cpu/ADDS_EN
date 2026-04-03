"""
ADDS Architecture Report Generator
Converts the architecture analysis to a professional DOCX document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_title_page(doc):
    """Add professional title page"""
    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ADDS 시스템\n아키텍처 분석 보고서')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(51, 51, 51)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('\nAI-Driven Anti-cancer Drug Development System\nComprehensive Architecture Analysis')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Add space
    for _ in range(5):
        doc.add_paragraph()
    
    # Metadata
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = [
        f'보고서 생성일: {datetime.now().strftime("%Y년 %m월 %d일")}',
        '',
        '발행 기관: 인하대학교 의생명공학과',
        '시스템 버전: 1.0',
        '분석 범위: 전체 코드베이스 (788,000+ 라인)',
    ]
    info_run = info.add_run('\n'.join(info_text))
    info_run.font.size = Pt(11)
    
    doc.add_page_break()

def add_heading_1(doc, text):
    """Add level 1 heading with custom style"""
    heading = doc.add_heading(text, level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(20)

def add_heading_2(doc, text):
    """Add level 2 heading"""
    heading = doc.add_heading(text, level=2)
    heading.runs[0].font.color.rgb = RGBColor(51, 102, 153)
    heading.runs[0].font.size = Pt(16)

def add_heading_3(doc, text):
    """Add level 3 heading"""
    heading = doc.add_heading(text, level=3)
    heading.runs[0].font.size = Pt(14)

def add_body_text(doc, text):
    """Add body paragraph with proper formatting"""
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_bullet_point(doc, text, level=0):
    """Add bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        run.font.size = Pt(11)

def add_code_block(doc, code, language='python'):
    """Add code block with monospace font"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 51, 51)
    
    # Add shading
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    p._element.get_or_add_pPr().append(shading_elm)

def create_architecture_report():
    """Generate the complete architecture report"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)
    
    # Title page
    add_title_page(doc)
    
    # Executive Summary
    add_heading_1(doc, '1. 개요 (Executive Summary)')
    
    add_body_text(doc, 
        'ADDS(AI-Driven Anti-cancer Drug Development System)는 인하대학교 의생명공학과에서 개발한 '
        'AI 기반 정밀 종양학 통합 플랫폼입니다. 본 시스템은 세포 이미지 분석, 약물 조합 최적화, '
        '임상 의사결정 지원 기능을 하나의 플랫폼에서 제공합니다.'
    )
    
    add_heading_2(doc, '1.1 시스템 목적')
    add_body_text(doc, 
        'ADDS는 다음 네 가지 핵심 기능을 통해 항암제 개발 프로세스를 혁신하고자 합니다:'
    )
    
    add_bullet_point(doc, '세포 이미지 자동 분석: Cellpose 기반 딥러닝 세그멘테이션으로 병리 슬라이드에서 세포를 자동으로 인식하고 정량화합니다.')
    add_bullet_point(doc, '약물 조합 최적화: Exscientia의 DTOL(Design-Test-Optimize-Learn) 사이클을 적용하여 최적의 항암제 조합을 탐색합니다.')
    add_bullet_point(doc, '임상 의사결정 지원: 환자의 유전자 변이, 바이오마커, 병리 데이터를 통합하여 개인 맞춤형 치료 계획을 제시합니다.')
    add_bullet_point(doc, '다중모달 데이터 통합: 이미지, 유전체, 임상 문서 등 다양한 데이터 소스를 하나의 플랫폼에서 관리하고 분석합니다.')
    
    add_heading_2(doc, '1.2 핵심 성과 지표')
    
    # Create table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '항목'
    hdr_cells[1].text = '수치'
    
    # Data rows
    data = [
        ('총 코드 라인 수', '788,000+ 라인'),
        ('모듈 수', '16개 주요 모듈'),
        ('클래스/구성 요소', '64+ 개'),
        ('GPU 성능 향상', '5배 (CPU 대비)')
    ]
    
    for i, (item, value) in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = item
        row_cells[1].text = value
    
    doc.add_paragraph()  # Space after table
    
    # Architecture Overview
    add_heading_1(doc, '2. 전체 아키텍처')
    
    add_heading_2(doc, '2.1 계층 구조')
    add_body_text(doc, 
        'ADDS는 4개의 계층으로 구성된 모듈형 아키텍처를 채택하고 있습니다. '
        '각 계층은 명확한 책임을 가지며, 하위 계층에만 의존하는 단방향 의존성 원칙을 따릅니다.'
    )
    
    add_heading_3(doc, 'UI Layer (사용자 인터페이스 계층)')
    add_body_text(doc, 
        'Streamlit 프레임워크 기반의 웹 인터페이스로, 사용자와 시스템 간의 상호작용을 담당합니다. '
        '14개의 페이지 모듈로 구성되어 있으며, 각 페이지는 특정 워크플로우를 지원합니다.'
    )
    
    add_bullet_point(doc, '홈 페이지: 시스템 개요 및 주요 메트릭 대시보드')
    add_bullet_point(doc, '이미지 분석: Cellpose 세그멘테이션, 심층 분석, 배치 처리 (54,405 라인)')
    add_bullet_point(doc, '정밀 종양학: 환자 등록, 병리 분석, 치료 추천 (65,017 라인)')
    add_bullet_point(doc, '배치 처리: 다중 이미지 일괄 분석 (10,979 라인)')
    add_bullet_point(doc, '문서 처리: PDF/DOCX AI 분석 (25,152 라인)')
    add_bullet_point(doc, '데이터 관리: 분석 결과 데이터베이스 관리')
    
    add_heading_3(doc, 'Application Layer (애플리케이션 계층)')
    add_body_text(doc, 
        '핵심 비즈니스 로직이 구현된 계층으로, 6개의 주요 모듈로 구성됩니다:'
    )
    
    add_bullet_point(doc, 'Preprocessing 모듈: 이미지 전처리 및 문서 파싱 (11개 파일)')
    add_bullet_point(doc, 'Recommendation 모듈: DTOL 엔진, Active Learning (13개 파일)')
    add_bullet_point(doc, 'Clinical 모듈: 환자 데이터베이스, 코호트 분류 (3개 파일)')
    add_bullet_point(doc, 'Pathology 모듈: 공간 분석, 이질성 메트릭 (4개 파일)')
    add_bullet_point(doc, 'Protein 모듈: ESMFold, PPI 네트워크 (4개 파일)')
    add_bullet_point(doc, 'AI 모듈: 파인튜닝, 데이터셋 빌더 (3개 파일)')
    
    add_heading_3(doc, 'Core Services Layer (핵심 서비스 계층)')
    add_body_text(doc, 
        '공통 유틸리티 및 인프라 서비스를 제공하는 계층입니다. 27개의 유틸리티 모듈이 포함되어 있으며, '
        '데이터베이스 관리, 캐싱, GPU 모니터링, 배치 처리 등의 기능을 담당합니다.'
    )
    
    add_heading_3(doc, 'Infrastructure Layer (인프라 계층)')
    add_body_text(doc, 
        '외부 라이브러리 및 하드웨어 자원을 관리하는 최하위 계층입니다. '
        'PyTorch, Cellpose, scikit-learn 등의 ML 프레임워크와 SQLite 데이터베이스, '
        'CUDA GPU 가속 등의 인프라 구성 요소가 포함됩니다.'
    )
    
    # Module Details
    add_heading_1(doc, '3. 모듈별 상세 분석')
    
    add_heading_2(doc, '3.1 UI 모듈 (src/ui/)')
    
    add_heading_3(doc, '개요')
    add_body_text(doc, 
        'UI 모듈은 Streamlit 프레임워크를 기반으로 한 웹 인터페이스로, 총 14개의 Python 파일과 '
        '4개의 서브디렉토리로 구성되어 있습니다. 전체 약 195,000 라인의 코드로 이루어져 있으며, '
        '사용자 친화적인 인터페이스를 통해 복잡한 분석 워크플로우를 간편하게 실행할 수 있도록 설계되었습니다.'
    )
    
    add_heading_3(doc, '주요 구성 요소')
    
    add_body_text(doc, '1) Main App (app.py, 255 라인)')
    add_body_text(doc, 
        '애플리케이션의 진입점(Entry Point)으로, 다음 기능을 담당합니다:'
    )
    add_bullet_point(doc, '사이드바 네비게이션: 16개 메뉴 항목을 계층적으로 구성', 1)
    add_bullet_point(doc, 'GPU 설정 토글: 실시간 GPU 메모리 모니터링 및 CPU/GPU 모드 전환', 1)
    add_bullet_point(doc, '캐시 관리: 적중률 표시 및 수동 캐시 초기화 기능', 1)
    add_bullet_point(doc, '페이지 라우팅: session_state 기반 동적 페이지 전환', 1)
    
    add_code_block(doc, '''# GPU 토글 구현 예시
use_gpu = st.checkbox(
    "🚀 GPU 가속 사용",
    value=st.session_state.get('use_gpu', False),
    help=f"GPU: {torch.cuda.get_device_name(0)}"
)

if use_gpu:
    st.success(f"✓ GPU 활성화\\n{torch.cuda.get_device_name(0)}")
    memory_info = gpu_monitor.get_memory_info()
    st.metric("GPU 메모리", 
              f"{memory_info['allocated_mb']} MB",
              delta=f"{memory_info['utilization_percent']}% 사용중")''')
    
    add_body_text(doc, '2) Core Utilities (app_core.py, 160 라인)')
    add_body_text(doc, 
        '공통 기능을 중앙 집중화하여 코드 중복을 방지하고 성능을 최적화합니다:'
    )
    
    add_bullet_point(doc, 'GPU 자동 구성: CUDA_VISIBLE_DEVICES 환경 변수 설정, NVIDIA GPU 강제 선택', 1)
    add_bullet_point(doc, '캐시된 리소스 팩토리: Cellpose 모델, Document Parser, Synergy Calculator 재사용', 1)
    add_bullet_point(doc, '전역 CSS 스타일: 일관된 UI 디자인 적용', 1)
    
    add_code_block(doc, '''@st.cache_resource(show_spinner=False)
def get_cellpose_processor(model_type: str, gpu: bool = False):
    """
    캐시된 Cellpose processor 반환
    모델은 첫 실행 시에만 로딩되고 이후에는 재사용
    """
    logger.info(f"Creating Cellpose: model={model_type}, gpu={gpu}")
    return CellposeProcessor(model_type=model_type, gpu=gpu)

# 효과: 모델 로딩 8초 → 0.1초 (재사용 시)''')
    
    # Continue with more sections...
    add_heading_2(doc, '3.2 전처리 모듈 (src/preprocessing/)')
    
    add_heading_3(doc, 'Cellpose 이미지 프로세서 (image_processor.py)')
    add_body_text(doc, 
        'CellposeProcessor 클래스는 604라인의 코드로 구성된 핵심 이미지 분석 엔진입니다. '
        'Cellpose 딥러닝 모델을 래핑하여 다음과 같은 기능을 제공합니다:'
    )
    
    add_bullet_point(doc, '자동 세포 세그멘테이션: cyto, cyto2, nuclei 모델 지원')
    add_bullet_point(doc, '형태학적 특징 추출: 19가지 정량 지표 계산 (면적, 둘레, 원형도, 고형도 등)')
    add_bullet_point(doc, '배치 처리: GPU/CPU 기반 다중 이미지 동시 처리')
    add_bullet_point(doc, '심층 분석: 세포 건강도, 밀집도, 이질성 평가')
    add_bullet_point(doc, '시각화 파이프라인: 6단계 이미지 생성 (원본→전처리→마스크→윤곽→히트맵→오버레이)')
    
    add_body_text(doc, '주요 메서드:')
    
    add_code_block(doc, '''class CellposeProcessor:
    def __init__(self, model_type='cyto2', gpu=True, batch_size=8):
        """초기화: 모델 로딩 및 GPU 설정"""
        
    def segment_image(self, image, diameter=None, **kwargs):
        """단일 이미지 세그멘테이션"""
        # Returns: (masks, flows, metadata)
        
    def extract_morphological_features(self, image, masks):
        """19가지 형태학적 특징 추출"""
        # Returns: DataFrame with cell-level features
        
    def calculate_cell_viability_metrics(self, features_df):
        """세포 건강도 및 집합 메트릭 계산"""
        # Returns: Dict with aggregate metrics
        
    def process_and_save(self, image_path, output_dir, deep_analysis=True):
        """종합 파이프라인: 세그멘테이션 + 특징 추출 + 저장"""''')
    
    add_body_text(doc, '성능 최적화:')
    add_bullet_point(doc, 'GPU 가속: RTX 5070 기준 이미지당 3초 (CPU 대비 5배 향상)', 1)
    add_bullet_point(doc, '배치 처리: 10개 이미지 18초 (순차 처리 60초 대비 70% 단축)', 1)
    add_bullet_point(doc, '메모리 효율화: 대용량 이미지 타일링 처리', 1)
    
    # DTOL Engine
    add_heading_2(doc, '3.3 DTOL 엔진 (src/recommendation/)')
    
    add_body_text(doc, 
        'DTOL(Design-Test-Optimize-Learn) 엔진은 Exscientia의 AI 기반 약물 발견 방법론을 '
        '항암제 조합 최적화에 적용한 모듈입니다. 13개 파일, 약 180,000 라인의 코드로 구성되며, '
        'Active Learning과 Multi-Parameter Optimization을 핵심 알고리즘으로 사용합니다.'
    )
    
    add_heading_3(doc, 'Combination Designer (조합 설계기)')
    
    add_body_text(doc, 
        'CombinationDesigner 클래스는 약물 조합을 지능적으로 설계하는 핵심 컴포넌트입니다. '
        '999 라인의 코드로 구성되어 있으며, 다음과 같은 알고리즘을 구현합니다:'
    )
    
    add_body_text(doc, '1) Multi-Parameter Optimization (MPO)')
    add_body_text(doc, 
        '약물 조합을 평가할 때 단일 지표가 아닌 5가지 파라미터를 종합적으로 고려합니다:'
    )
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '파라미터'
    hdr_cells[1].text = '가중치'
    hdr_cells[2].text = '설명'
    
    # Data
    mpo_data = [
        ('Safety (안전성)', '30%', '독성, 약물 상호작용, 신기능/간기능 적합성'),
        ('Efficacy (효능)', '35%', '예상 종양 감소율, 바이오마커 반응'),
        ('Synergy (시너지)', '20%', '약물 간 상승 효과, 경로 중복도'),
        ('Cost (비용)', '10%', '치료 비용, 보험 적용 여부'),
        ('Feasibility (실현성)', '5%', '투여 편의성, 환자 순응도')
    ]

    for i, (param, weight, desc) in enumerate(mpo_data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = param
        row_cells[1].text = weight
        row_cells[2].text = desc
    
    doc.add_paragraph()
    
    add_body_text(doc, '2) Active Learning (능동 학습)')
    add_body_text(doc, 
        'Gaussian Process를 사용한 베이지안 최적화로, 실험 횟수를 최소화하면서 최적 조합을 탐색합니다. '
        '초기 10회는 탐색(Exploration) 위주, 이후에는 활용(Exploitation) 위주로 전환하는 Dual Mode를 채택했습니다.'
    )
    
    add_bullet_point(doc, 'Expected Improvement (EI): 현재 최선 대비 개선 기대값', 1)
    add_bullet_point(doc, 'Upper Confidence Bound (UCB): 불확실성과 평균의 균형', 1)
    add_bullet_point(doc, 'Thompson Sampling: 확률적 샘플링', 1)
    add_bullet_point(doc, 'Entropy Search: 정보 획득 최대화', 1)
    
    # Medical Safety
    add_heading_1(doc, '4. 의료 안전성 및 규제 준수')
    
    add_body_text(doc, 
        'ADDS는 의료기기 및 임상시험 규제를 엄격히 준수하도록 설계되었습니다. '
        '모든 약물 추천 및 독성 평가는 국제 표준을 따르며, 환자 안전을 최우선으로 합니다.'
    )
    
    add_heading_2(doc, '4.1 CTCAE v5.0 독성 평가')
    
    add_body_text(doc, 
        'Common Terminology Criteria for Adverse Events (CTCAE) 버전 5.0을 기준으로 '
        '약물 조합의 독성을 정량화합니다. 4가지 장기 시스템별로 Grade 0-4 등급을 부여하며, '
        'Grade 4 혈액독성 또는 Grade 3-4 비혈액독성이 발생하면 DLT(Dose-Limiting Toxicity)로 자동 판정합니다.'
    )
    
    add_code_block(doc, '''class ToxicityProfile:
    """CTCAE v5.0 기준 독성 프로필"""
    
    hematologic: Dict[str, int]      # 백혈구 감소, 빈혈 등
    hepatic: Dict[str, int]          # ALT/AST 상승, 황달
    renal: Dict[str, int]            # 크레아티닌 상승
    gastrointestinal: Dict[str, int] # 오심, 구토, 설사
    
    def is_dose_limiting_toxicity(self) -> bool:
        """DLT 판정"""
        max_grade = self.calculate_max_grade()
        
        # Grade 4 hematologic OR Grade 3-4 non-hematologic
        if max_grade >= 4:
            return True
        if max(self.hepatic.values(), 
               self.renal.values(),
               self.gastrointestinal.values()) >= 3:
            return True
        
        return False''')
    
    add_heading_2(doc, '4.2 FDA 용량 규제 준수')
    
    add_body_text(doc, 
        '모든 약물은 FDA 승인 용량 범위 내에서만 사용됩니다. Drug 클래스는 초기화 시 '
        'min_dose, max_dose, standard_dose를 검증하며, 범위를 벗어난 용량은 자동으로 거부됩니다.'
    )
    
    add_heading_2(doc, '4.3 HIPAA 개인정보 보호')
    
    add_body_text(doc, 
        '환자 데이터는 Health Insurance Portability and Accountability Act (HIPAA) 기준을 준수합니다. '
        '식별 가능한 개인정보(이름, 주민번호, 주소 등)는 저장하지 않으며, 익명화된 환자 ID와 '
        '최소한의 임상 정보(연령, 병기)만 데이터베이스에 기록합니다.'
    )
    
    # Performance
    add_heading_1(doc, '5. 성능 최적화 및 벤치마크')
    
    add_heading_2(doc, '5.1 GPU 가속')
    
    add_body_text(doc, 
        'NVIDIA RTX 5070 GPU를 활용하여 Cellpose 세그멘테이션 속도를 5배 향상시켰습니다.'
    )
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '작업'
    hdr_cells[1].text = 'CPU 시간'
    hdr_cells[2].text = 'GPU 시간'
    
    perf_data = [
        ('단일 이미지 세그멘테이션', '15초', '3초 (5배 향상)'),
        ('배치 10개 이미지', '150초', '30초'),
        ('모델 로딩 (최초)', '8초', '8초')
    ]
    
    for i, (task, cpu, gpu) in enumerate(perf_data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = task
        row_cells[1].text = cpu
        row_cells[2].text = gpu
    
    doc.add_paragraph()
    
    add_heading_2(doc, '5.2 캐싱 전략')
    
    add_body_text(doc, 
        'Streamlit의 @st.cache_resource 데코레이터를 활용하여 무거운 객체를 재사용합니다. '
        '현재 시스템 운영 중 캐시 적중률은 78.9%에 달하며, 모델 로딩 시간을 99% 단축시켰습니다.'
    )
    
    # Conclusion
    add_heading_1(doc, '6. 결론 및 향후 계획')
    
    add_body_text(doc, 
        'ADDS는 788,000 라인의 코드와 16개 모듈로 구성된 종합 정밀 종양학 플랫폼입니다. '
        'Cellpose 딥러닝 세그멘테이션, Exscientia DTOL 사이클, PK/PD 시뮬레이션, '
        '다중모달 데이터 통합 등 최첨단 AI/ML 기술을 의료 안전성 기준과 결합하여 구현했습니다.'
    )
    
    add_heading_2(doc, '6.1 핵심 성과')
    
    add_bullet_point(doc, '의료 안전성: CTCAE v5.0, FDA, HIPAA 준수')
    add_bullet_point(doc, 'AI/ML 통합: Cellpose, GPT-4, Gaussian Process, ESMFold')
    add_bullet_point(doc, '성능: GPU 5배 향상, 캐싱 79% 적중률, 병렬 처리 70% 시간 단축')
    add_bullet_point(doc, '사용성: 직관적 Streamlit UI, 실시간 피드백, 한글 지원')
    
    add_heading_2(doc, '6.2 차별화 포인트')
    
    add_body_text(doc, 
        'ADDS는 국내 최초로 Exscientia의 DTOL 사이클을 항암제 조합 최적화에 적용했습니다. '
        '또한 PK/PD 시뮬레이션을 통한 독성 예측, 다중모달 데이터 통합(이미지+유전자+임상), '
        'Active Learning 기반 실험 횟수 최소화 등 연구 수준을 넘어 실제 임상 의사결정 지원이 '
        '가능한 수준의 시스템을 구축했습니다.'
    )
    
    add_heading_2(doc, '6.3 향후 개선 계획')
    
    add_body_text(doc, '단기 (1개월):')
    add_bullet_point(doc, 'Legacy UI 파일 정리 및 통합', 1)
    add_bullet_point(doc, '단위 테스트 커버리지 80% 달성', 1)
    add_bullet_point(doc, 'Docker 배포 자동화', 1)
    
    add_body_text(doc, '중기 (3개월):')
    add_bullet_point(doc, '실제 임상 데이터 통합 및 검증', 1)
    add_bullet_point(doc, 'Multi-GPU 분산 처리 지원', 1)
    add_bullet_point(doc, 'REST API 개발 (FastAPI)', 1)
    
    add_body_text(doc, '장기 (6개월+):')
    add_bullet_point(doc, 'Federated Learning을 통한 다기관 협업', 1)
    add_bullet_point(doc, '실시간 예측 (< 1초 응답)', 1)
    add_bullet_point(doc, '클라우드 배포 (AWS/GCP)', 1)
    
    # Final page
    doc.add_page_break()
    
    # Appendix
    add_heading_1(doc, '부록 A: 기술 스택 상세')
    
    table = doc.add_table(rows=11, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '카테고리'
    hdr_cells[1].text = '기술'
    hdr_cells[2].text = '버전/용도'
    
    tech_data = [
        ('Deep Learning', 'PyTorch', '2.x, GPU 가속 및 모델 학습'),
        ('Computer Vision', 'Cellpose', '2.x, 세포 세그멘테이션'),
        ('Machine Learning', 'scikit-learn', '1.x, 클러스터링/회귀'),
        ('Optimization', 'SciPy', '1.x, Gaussian Process'),
        ('Data Processing', 'Pandas/NumPy', 'Latest, 데이터 조작'),
        ('Database', 'SQLite', '3.x, 관계형 DB'),
        ('UI Framework', 'Streamlit', '1.30+, 웹 UI'),
        ('Visualization', 'Plotly', '인터랙티브 차트'),
        ('Container', 'Docker', '배포 환경'),
        ('GPU', 'CUDA', '12.1, GPU 가속')
    ]
    
    for i, (cat, tech, ver) in enumerate(tech_data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = cat
        row_cells[1].text = tech
        row_cells[2].text = ver
    
    doc.add_paragraph()
    
    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n\n보고서 생성일: {datetime.now().strftime("%Y년 %m월 %d일")}\n인하대학교 의생명공학과')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    output_path = 'C:/Users/brook/Desktop/ADDS_Architecture_Report.docx'
    doc.save(output_path)
    print(f"Report saved: {output_path}")
    return output_path

if __name__ == '__main__':
    output_path = create_architecture_report()
    print(f"Report saved successfully: {output_path}")
